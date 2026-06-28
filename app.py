import os
import streamlit as st
from docx import Document
from io import BytesIO
import subprocess

# ==========================================
# ORIGINAL SCRIPT LOGIC
# ==========================================

def replace_placeholders(doc, data_map):
    """
    Scans a Word document's paragraphs and tables to replace 
    placeholders while preserving the original formatting.
    """
    for paragraph in doc.paragraphs:
        for placeholder, new_text in data_map.items():
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(new_text))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, new_text in data_map.items():
                        if placeholder in paragraph.text:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(new_text))

def convert_docx_to_pdf(docx_bytes):
    """
    Converts DOCX bytes to PDF using LibreOffice (supported natively on Streamlit Cloud Linux servers)
    """
    # Write the docx bytes to a temporary file
    with open("temp_input.docx", "wb") as f:
        f.write(docx_bytes)
    
    # Run headless LibreOffice to convert docx to pdf
    cmd = ["libreoffice", "--headless", "--convert-to", "pdf", "temp_input.docx"]
    subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    
    # Read the generated PDF back into memory
    if os.path.exists("temp_input.pdf"):
        with open("temp_input.pdf", "rb") as f:
            pdf_bytes = f.read()
        
        # Clean up temporary files
        os.remove("temp_input.docx")
        os.remove("temp_input.pdf")
        return pdf_bytes
    else:
        raise Exception("PDF conversion failed. Check server dependencies.")

# ==========================================
# STREAMLIT USER INTERFACE
# ==========================================

st.set_page_config(page_title="AI CV & Document Generator", page_icon="📄", layout="centered")

st.title("📄 Multi-Format CV Generator")
st.write("Upload your styled Word template (`.docx`), fill in your details, and choose your preferred output format!")

st.markdown("---")

# Step 1: File Upload (Still uses .docx as the base engine)
st.subheader("1. Upload Your Word Template")
uploaded_file = st.file_uploader("Choose a template file (.docx)", type=["docx"])

with st.expander("💡 Why only .docx templates?"):
    st.write("PDFs are uneditable layouts. By uploading a `.docx` template, the app can intelligently swap text while preserving your design, then export it beautifully to PDF for you!")

st.markdown("---")

# Step 2: Form Input Fields
st.subheader("2. Enter Your New Information")

with st.form("cv_data_form"):
    col1, col2 = st.columns(2)
    with col1:
        full_name = st.text_input("Full Name", "Alex Mercer")
        job_title = st.text_input("Job Title", "Senior AI Prompt Engineer")
        phone = st.text_input("Phone Number", "+1 (555) 019-2834")
    with col2:
        email = st.text_input("Email Address", "alex.mercer@email.com")
        linkedin = st.text_input("LinkedIn Profile", "linkedin.com/in/alexmercer")
        education = st.text_input("Education", "B.S. in Computer Science - Stanford University")
        
    summary = st.text_area("Professional Summary", "Results-driven AI Professional with experience specializing in LLM deployment.")
    
    st.markdown("**Work Experience Section**")
    exp_title = st.text_input("Job 1 Title & Company", "Lead AI Engineer | TechCorp")
    exp_desc = st.text_area("Job 1 Description", "• Designed and deployed AI workflows.\n• Reduced time-to-hire by 40%.")
    
    st.markdown("---")
    st.markdown("**Export Settings**")
    output_format = st.radio("Select Output Format:", ["Word Document (.docx)", "Adobe PDF (.pdf)"])
    
    submit_button = st.form_submit_button(label="⚡ Generate Document")

# Step 3: Process and Download
if submit_button:
    if uploaded_file is not None:
        NEW_CV_DATA = {
            "{{FULL_NAME}}": full_name,
            "{{JOB_TITLE}}": job_title,
            "{{PHONE}}": phone,
            "{{EMAIL}}": email,
            "{{LINKEDIN}}": linkedin,
            "{{PROFESSIONAL_SUMMARY}}": summary,
            "{{EXP_JOB_1}}": exp_title,
            "{{EXP_DESC_1}}": exp_desc,
            "{{EDUCATION}}": education
        }
        
        try:
            # Process Word Document
            doc = Document(uploaded_file)
            replace_placeholders(doc, NEW_CV_DATA)
            
            docx_stream = BytesIO()
            doc.save(docx_stream)
            docx_bytes = docx_stream.getvalue()
            
            st.success("🎉 Success! Your document has been built.")
            
            if output_format == "Word Document (.docx)":
                st.download_button(
                    label="📥 Download Word File (.docx)",
                    data=docx_bytes,
                    file_name="generated_cv.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                with st.spinner("Converting document to PDF..."):
                    pdf_bytes = convert_docx_to_pdf(docx_bytes)
                
                st.download_button(
                    label="📥 Download PDF File (.pdf)",
                    data=pdf_bytes,
                    file_name="generated_cv.pdf",
                    mime="application/pdf"
                )
                
        except Exception as e:
            st.error(f"An error occurred: {e}")
    else:
        st.warning("⚠️ Please upload a template `.docx` file first.")
